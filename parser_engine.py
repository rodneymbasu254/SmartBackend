"""
parser_api.py

Simple FastAPI parser microservice for SmartLearningAI.
- Accepts .docx and .pdf course outlines
- Extracts structured data (course name, code, lecturer, weekly topics, etc.)
- Returns clean JSON
"""

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from pathlib import Path
import tempfile
import json
import fitz  # PyMuPDF for PDF support


app = FastAPI(title="SmartLearningAI Parser", version="1.0")

# Allow all origins for now (you can restrict later)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --------------------------
# DOCX PARSER
# --------------------------
def parse_docx_file(path):
    doc = Document(path)

    data = {
        "course_code": None,
        "course_name": None,
        "prerequisite": None,
        "credit_hours": None,
        "lecturer": None,
        "email": None,
        "purpose": None,
        "objectives": None,
        "weekly_topics": [],
        "assessment": None,
        "core_reading": None,
        "references": None
    }

    # --- Extract from tables ---
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                val = row.cells[1].text.strip()

                if "course code" in key:
                    data["course_code"] = val
                elif "course name" in key:
                    data["course_name"] = val
                elif "prerequisite" in key:
                    data["prerequisite"] = val
                elif "credit" in key:
                    data["credit_hours"] = val
                elif "lecturer" in key:
                    data["lecturer"] = val
                elif "email" in key:
                    data["email"] = val
                elif key.startswith("week"):
                    data["weekly_topics"].append(f"{row.cells[0].text.strip()} {val}")

    # --- Extract from paragraphs ---
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for i, p in enumerate(paragraphs):
        lower = p.lower()

        if "purpose of the course" in lower and i + 1 < len(paragraphs):
            data["purpose"] = paragraphs[i + 1]

        elif "objective" in lower and i + 1 < len(paragraphs):
            data["objectives"] = paragraphs[i + 1]

        elif "course assessment" in lower:
            collect = []
            j = i + 1
            while j < len(paragraphs) and "core reading" not in paragraphs[j].lower():
                collect.append(paragraphs[j])
                j += 1
            data["assessment"] = "\n".join(collect)

        elif "core reading" in lower and i + 1 < len(paragraphs):
            data["core_reading"] = paragraphs[i + 1]

        elif "recommended reference" in lower and i + 1 < len(paragraphs):
            data["references"] = paragraphs[i + 1]

    return data


# --------------------------
# PDF PARSER
# --------------------------
def parse_pdf_file(path):
    """Convert PDF text into a temporary DOCX-like structure."""
    doc = fitz.open(path)
    text_lines = []
    for page in doc:
        text = page.get_text("text")
        if text:
            text_lines.extend([t.strip() for t in text.splitlines() if t.strip()])
    doc.close()

    # simple heuristic: simulate the same logic used for docx
    data = {
        "course_code": None,
        "course_name": None,
        "prerequisite": None,
        "credit_hours": None,
        "lecturer": None,
        "email": None,
        "purpose": None,
        "objectives": None,
        "weekly_topics": [],
        "assessment": None,
        "core_reading": None,
        "references": None
    }

    for i, line in enumerate(text_lines):
        lower = line.lower()

        if "course code" in lower:
            data["course_code"] = line.split(":")[-1].strip()
        elif "course name" in lower:
            data["course_name"] = line.split(":")[-1].strip()
        elif "prerequisite" in lower:
            data["prerequisite"] = line.split(":")[-1].strip()
        elif "credit" in lower:
            data["credit_hours"] = line.split(":")[-1].strip()
        elif "lecturer" in lower:
            data["lecturer"] = line.split(":")[-1].strip()
        elif "email" in lower:
            data["email"] = line.split(":")[-1].strip()
        elif "purpose of the course" in lower and i + 1 < len(text_lines):
            data["purpose"] = text_lines[i + 1]
        elif "objective" in lower and i + 1 < len(text_lines):
            data["objectives"] = text_lines[i + 1]
        elif lower.startswith("week"):
            data["weekly_topics"].append(line)
        elif "course assessment" in lower:
            collect = []
            j = i + 1
            while j < len(text_lines) and "core reading" not in text_lines[j].lower():
                collect.append(text_lines[j])
                j += 1
            data["assessment"] = "\n".join(collect)
        elif "core reading" in lower and i + 1 < len(text_lines):
            data["core_reading"] = text_lines[i + 1]
        elif "recommended reference" in lower and i + 1 < len(text_lines):
            data["references"] = text_lines[i + 1]

    return data


# --------------------------
# API ENDPOINTS
# --------------------------
@app.post("/api/parse")
async def parse_outline(file: UploadFile = File(...)):
    """Accepts .docx or .pdf and returns structured JSON."""
    ext = Path(file.filename).suffix.lower()

    if ext not in [".docx", ".pdf"]:
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files supported")

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            contents = await file.read()
            tmp.write(contents)
            tmp.flush()
            path = tmp.name

        if ext == ".docx":
            parsed = parse_docx_file(path)
        else:
            parsed = parse_pdf_file(path)

        return JSONResponse(content={"status": "ok", "data": parsed})

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")
    finally:
        try:
            Path(path).unlink(missing_ok=True)
        except Exception:
            pass


@app.get("/health")
async def health():
    """Simple health check endpoint."""
    return {"status": "ok", "service": "parser", "version": "1.0"}
